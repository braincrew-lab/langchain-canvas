"""Word editing — the invariants an in-place document edit has to keep.

The load-bearing one is the first: whatever the operation did NOT touch comes
out of the ZIP byte for byte. A Word package that is re-serialized whole picks
up drift nobody asked for — direct formatting multiplies, font references get
rewritten — and the user sees their document come back subtly different from
the one they uploaded. Every operation here is measured against that.

The rest guard the ways an edit can be silently wrong: an address that means a
different paragraph on the second read, an anchor that matched nothing and was
shrugged off, a merged cell counted three times, a picture swap that stretched
the page.
"""

from __future__ import annotations

import hashlib
import io
import zipfile

import pytest
from documents import FIXTURES, png_bytes, sample_document

from langchain_canvas import document_ops as ops


def _entries(data: bytes) -> dict[str, str]:
    with zipfile.ZipFile(io.BytesIO(data)) as archive:
        return {name: hashlib.sha256(archive.read(name)).hexdigest() for name in archive.namelist()}


def _changed(before: bytes, after: bytes) -> set[str]:
    old, new = _entries(before), _entries(after)
    return {name for name in new if old.get(name) != new[name]}


@pytest.fixture
def document() -> bytes:
    return sample_document()


# --- I1: parts the edit did not touch travel through unchanged -------------------


def test_text_replacement_touches_only_the_body_part(document: bytes) -> None:
    after = ops.replace_text(document, "즉시 조치가 필요한", "즉시 조치가 반드시 필요한")
    assert _changed(document, after) == {"word/document.xml"}


def test_header_replacement_touches_only_the_header_part(document: bytes) -> None:
    after = ops.replace_text(document, "Braincrew Confidential", "Braincrew Internal")
    assert _changed(document, after) == {"word/header1.xml"}


def test_insert_touches_only_the_body_part(document: bytes) -> None:
    after = ops.insert_paragraph(document, anchor="사진 1. 점검 당일 현장", text="사진 2. 보완 후")
    assert _changed(document, after) == {"word/document.xml"}


def test_remove_touches_only_the_body_part(document: bytes) -> None:
    after = ops.remove_paragraph(document, anchor="사진 1. 점검 당일 현장")
    assert _changed(document, after) == {"word/document.xml"}


def test_image_replacement_touches_only_its_own_parts(document: bytes) -> None:
    after, _ = ops.replace_image(document, index=0, image=png_bytes(300, 60, (0x8B, 0, 0)))
    changed = _changed(document, after)
    assert changed <= {
        "word/document.xml",
        "word/_rels/document.xml.rels",
        "[Content_Types].xml",
        "word/media/image2.png",
    }
    assert "word/document.xml" in changed
    assert "word/styles.xml" not in changed


def test_styles_and_theme_survive_every_operation(document: bytes) -> None:
    """The parts that decide how the document *looks* are never rewritten."""
    untouched = {"word/styles.xml", "word/settings.xml", "word/theme/theme1.xml"}
    results = [
        ops.replace_text(document, "즉시 조치", "즉각 조치"),
        ops.insert_paragraph(document, anchor="사진 1. 점검 당일 현장", text="추가"),
        ops.remove_paragraph(document, anchor="사진 1. 점검 당일 현장"),
        ops.replace_image(document, index=0, image=png_bytes(40, 40, (1, 2, 3)))[0],
    ]
    before = _entries(document)
    for after in results:
        now = _entries(after)
        for name in untouched & set(before):
            assert now[name] == before[name], name


def test_repack_keeps_entry_order_and_compression(document: bytes) -> None:
    after = ops.replace_text(document, "즉시 조치", "즉각 조치")
    with zipfile.ZipFile(io.BytesIO(document)) as before, zipfile.ZipFile(io.BytesIO(after)) as now:
        assert now.namelist() == before.namelist()
        kinds_before = {i.filename: i.compress_type for i in before.infolist()}
        kinds_now = {i.filename: i.compress_type for i in now.infolist()}
        assert kinds_now == kinds_before


# --- I2: an address means the same paragraph on every read ----------------------


def test_paragraph_addresses_are_stable_across_reads(document: bytes) -> None:
    first = ops.outline(document)
    second = ops.outline(document)
    assert first.lines == second.lines
    assert first.counts == second.counts


def test_addresses_index_document_paragraphs_including_blank_ones(document: bytes) -> None:
    """``[pN]`` is ``document.paragraphs[N]`` — the number never skips a blank."""
    from docx import Document

    parsed = Document(io.BytesIO(document))
    shown = {
        int(line[2 : line.index("]")]): line
        for line in ops.outline(document).lines
        if line.startswith("[p") and "]" in line
    }
    assert shown, "no paragraph lines rendered"
    for index, line in shown.items():
        text = parsed.paragraphs[index].text.strip()
        assert text and text in line
    blanks = [i for i, p in enumerate(parsed.paragraphs) if not p.text.strip()]
    assert blanks, "the fixture is meant to carry blank paragraphs"
    assert not (set(blanks) & set(shown)), "blank paragraphs are numbered, not printed"


def test_outline_names_styles_tables_pictures_and_counts(document: bytes) -> None:
    rendered = ops.outline(document).render()
    assert "paragraphs: 14" in rendered
    assert "tables: 2" in rendered and "images: 1" in rendered and "sections: 1" in rendered
    assert "[p0] (Heading 1) 2026 반영계획안" in rendered
    assert "(List Bullet)" in rendered
    assert "[t0] 3x3 table" in rendered
    assert "[img0] 2.4 x 1.6 in, in p8" in rendered
    assert "[header/p0] 브레인크루 | Braincrew Confidential" in rendered


# --- I3: an anchor that is not exactly one place fails loudly -------------------


def test_missing_anchor_names_the_closest_paragraph_and_the_difference(
    document: bytes,
) -> None:
    with pytest.raises(ops.AnchorError) as caught:
        ops.replace_text(document, "본 자료는 정보 제공 목적으로만은 작성되었으며", "…")
    message = str(caught.value)
    assert "0 matches" in message
    assert "Closest paragraph [p13]" in message
    assert "본 자료는 정보 제공 목적으로만 작성되었으며" in message
    assert "First difference at character 18" in message


def test_missing_anchor_that_shares_no_prefix_still_points_somewhere(
    document: bytes,
) -> None:
    with pytest.raises(ops.AnchorError) as caught:
        ops.replace_text(document, "존재하지 않는 문장", "…")
    assert "Closest paragraph [p" in str(caught.value)


def test_anchor_longer_than_the_paragraph_says_where_it_ends(document: bytes) -> None:
    with pytest.raises(ops.AnchorError) as caught:
        ops.remove_paragraph(
            document,
            anchor="본 자료는 정보 제공 목적으로만 작성되었으며 대외 배포를 금합니다. 끝.",
        )
    assert "ends after" in str(caught.value)


def test_repeated_anchor_is_refused_with_a_count_and_places(document: bytes) -> None:
    with pytest.raises(ops.AnchorError) as caught:
        ops.replace_text(document, "진행", "완료")
    message = str(caught.value)
    assert "matches 2 places" in message
    assert "[t1/r1c2/p0]" in message and "[t1/r2c2/p0]" in message


def test_a_merged_cell_is_one_place_not_three(document: bytes) -> None:
    """python-docx reports a merged cell once per grid column it spans."""
    after = ops.replace_text(document, "9월 조치 계획", "9월 조치 계획 (확정)")
    assert ops.reopens(after) is None


def test_multi_line_anchor_says_anchors_live_inside_a_paragraph(document: bytes) -> None:
    with pytest.raises(ops.AnchorError) as caught:
        ops.replace_text(document, "2026 반영계획안\n본 문서는", "…")
    assert "one paragraph" in str(caught.value)


def test_empty_anchor_is_refused(document: bytes) -> None:
    with pytest.raises(ops.AnchorError):
        ops.replace_text(document, "", "x")


# --- I7: the result is still a document ----------------------------------------


@pytest.mark.parametrize(
    "operation",
    [
        lambda d: ops.replace_text(d, "즉시 조치", "즉각 조치"),
        lambda d: ops.insert_paragraph(d, anchor="사진 1. 점검 당일 현장", text="새 문단"),
        lambda d: ops.remove_paragraph(d, anchor="사진 1. 점검 당일 현장"),
        lambda d: ops.replace_image(d, index=0, image=png_bytes(60, 30, (9, 9, 9)))[0],
    ],
)
def test_every_operation_reopens(document: bytes, operation) -> None:
    assert ops.reopens(operation(document)) is None


def test_reopens_reports_why_bad_bytes_are_not_a_document() -> None:
    assert ops.reopens(b"not a docx at all") is not None


# --- the edits themselves --------------------------------------------------------


def test_replacement_crosses_run_boundaries_and_keeps_the_first_run_format(
    document: bytes,
) -> None:
    from docx import Document

    after = ops.replace_text(document, "검토 결과, 즉시 조치가", "확인 결과, 즉시 조치가")
    parsed = Document(io.BytesIO(after))
    paragraph = parsed.paragraphs[4]
    assert paragraph.text == "확인 결과, 즉시 조치가 필요한 항목이 확인되었습니다."
    assert paragraph.runs[0].font.size is not None


def test_replacement_reaches_text_inside_a_hyperlink(document: bytes) -> None:
    after = ops.replace_text(document, "사내 포털", "사내 인트라넷")
    assert "사내 인트라넷" in ops.outline(after).render()


def test_replacement_reaches_a_table_cell(document: bytes) -> None:
    after = ops.replace_text(document, "소화기 위치 재배치", "소화기 위치 전면 재배치")
    assert "소화기 위치 전면 재배치" in ops.outline(after).render()


def test_insert_lands_after_the_anchor_by_default(document: bytes) -> None:
    from docx import Document

    after = ops.insert_paragraph(document, anchor="사진 1. 점검 당일 현장", text="사진 2. 보완 후")
    texts = [p.text for p in Document(io.BytesIO(after)).paragraphs]
    assert texts.index("사진 2. 보완 후") == texts.index("사진 1. 점검 당일 현장") + 1


def test_insert_can_land_before_the_anchor(document: bytes) -> None:
    from docx import Document

    after = ops.insert_paragraph(
        document, anchor="사진 1. 점검 당일 현장", text="머리말", position="before"
    )
    texts = [p.text for p in Document(io.BytesIO(after)).paragraphs]
    assert texts.index("머리말") == texts.index("사진 1. 점검 당일 현장") - 1


def test_insert_applies_a_named_style(document: bytes) -> None:
    from docx import Document

    after = ops.insert_paragraph(
        document, anchor="사진 1. 점검 당일 현장", text="4. 후속 조치", style="Heading 2"
    )
    styles = {p.text: p.style.name for p in Document(io.BytesIO(after)).paragraphs}
    assert styles["4. 후속 조치"] == "Heading 2"


def test_unknown_style_lists_the_styles_the_document_has(document: bytes) -> None:
    with pytest.raises(ops.DocumentOpError) as caught:
        ops.insert_paragraph(
            document, anchor="사진 1. 점검 당일 현장", text="x", style="Corporate Body"
        )
    message = str(caught.value)
    assert "no paragraph style named 'Corporate Body'" in message
    assert "Heading 1" in message


def test_bad_position_is_refused(document: bytes) -> None:
    with pytest.raises(ops.DocumentOpError):
        ops.insert_paragraph(document, anchor="사진 1. 점검 당일 현장", text="x", position="above")


def test_removing_the_only_paragraph_of_a_cell_is_refused(document: bytes) -> None:
    with pytest.raises(ops.DocumentOpError) as caught:
        ops.remove_paragraph(document, anchor="소화기 위치 재배치")
    assert "only paragraph in its container" in str(caught.value)


def test_remove_drops_exactly_one_paragraph(document: bytes) -> None:
    before = ops.outline(document).counts["paragraphs"]
    after = ops.remove_paragraph(document, anchor="사진 1. 점검 당일 현장")
    assert ops.outline(after).counts["paragraphs"] == before - 1


def test_image_replacement_keeps_the_width_and_refits_the_height(document: bytes) -> None:
    after, note = ops.replace_image(document, index=0, image=png_bytes(300, 60, (1, 2, 3)))
    from docx import Document

    shape = Document(io.BytesIO(after)).inline_shapes[0]
    assert round(shape.width.inches, 2) == 2.4
    assert round(shape.height.inches, 2) == round(2.4 * 60 / 300, 2)
    assert "width kept, height refitted" in note


def _media(data: bytes) -> list[str]:
    with zipfile.ZipFile(io.BytesIO(data)) as archive:
        return sorted(n for n in archive.namelist() if n.startswith("word/media/"))


def _with_picture(*, in_header: bool = False, twice: bool = False) -> bytes:
    """A document whose one picture is shown in the places named."""
    from docx import Document
    from docx.shared import Inches

    picture = png_bytes(120, 80, (0x2E, 0x7D, 0x32))
    document = Document()
    document.add_paragraph("intro")
    document.add_picture(io.BytesIO(picture), width=Inches(2.4))
    if twice:
        document.add_picture(io.BytesIO(picture), width=Inches(1.2))
    if in_header:
        run = document.sections[0].header.paragraphs[0].add_run()
        run.add_picture(io.BytesIO(picture), width=Inches(1.0))
    out = io.BytesIO()
    document.save(out)
    return out.getvalue()


def test_a_replaced_picture_takes_its_old_bytes_with_it() -> None:
    before = _with_picture()
    after, _ = ops.replace_image(before, index=0, image=png_bytes(60, 60, (0xC0, 0, 0)))
    assert _media(before) == ["word/media/image1.png"]
    assert _media(after) == ["word/media/image2.png"]
    assert ops.reopens(after) is None


def test_replacing_the_same_picture_again_does_not_grow_the_file() -> None:
    data = _with_picture()
    for shade in range(5):
        data, _ = ops.replace_image(data, index=0, image=png_bytes(60, 60, (shade, 9, 9)))
    assert len(_media(data)) == 1
    with zipfile.ZipFile(io.BytesIO(data)) as archive:
        assert len(archive.namelist()) == 18


def test_a_picture_the_header_also_shows_is_kept() -> None:
    """Only what nothing points at goes — the header's own link still reaches it."""
    before = _with_picture(in_header=True)
    after, _ = ops.replace_image(before, index=0, image=png_bytes(60, 60, (0xC0, 0, 0)))
    assert _media(after) == ["word/media/image1.png", "word/media/image2.png"]
    assert ops.reopens(after) is None


def test_a_picture_shown_twice_in_the_body_is_kept() -> None:
    before = _with_picture(twice=True)
    after, _ = ops.replace_image(before, index=0, image=png_bytes(60, 60, (0xC0, 0, 0)))
    assert "word/media/image1.png" in _media(after)
    from docx import Document

    assert len(Document(io.BytesIO(after)).inline_shapes) == 2


def test_dropping_a_picture_leaves_every_other_part_byte_identical() -> None:
    before = _with_picture()
    after, _ = ops.replace_image(before, index=0, image=png_bytes(60, 60, (0xC0, 0, 0)))
    old, new = _entries(before), _entries(after)
    edited = {
        "word/document.xml",
        "word/_rels/document.xml.rels",
        "[Content_Types].xml",
        "word/media/image1.png",
    }
    survivors = [name for name in old if name not in edited]
    assert survivors and all(new.get(name) == old[name] for name in survivors)


def test_repack_leaves_out_only_what_the_caller_named(document: bytes) -> None:
    """A part goes missing because an operation said so, never on its own."""
    saved = ops.repack(document, document, set())
    assert _entries(saved).keys() == _entries(document).keys()
    trimmed = ops.repack(document, document, set(), removed={"word/styles.xml"})
    assert "word/styles.xml" not in _entries(trimmed)


def test_image_index_out_of_range_names_the_range(document: bytes) -> None:
    with pytest.raises(ops.DocumentOpError) as caught:
        ops.replace_image(document, index=4, image=png_bytes(10, 10, (0, 0, 0)))
    assert "[img0]..[img0]" in str(caught.value)


def test_replacement_that_is_not_an_image_is_refused(document: bytes) -> None:
    with pytest.raises(ops.DocumentOpError) as caught:
        ops.replace_image(document, index=0, image=b"plain text")
    assert "not a readable image" in str(caught.value)


def test_unreadable_bytes_are_refused_before_anything_else() -> None:
    with pytest.raises(ops.DocumentPartError):
        ops.outline(b"PK\x03\x04 not really")


# --- I6: what the surface advertises is what the surface does -------------------


def test_every_advertised_format_has_a_document_that_proves_it() -> None:
    """A suffix in the tuple is a promise to the model; each one is exercised.

    Adding a format is cheap and silent — a tuple grows, every description
    picks the new suffix up, and nothing proves the parser behind it exists.
    This is the proof, one file per format.
    """
    for suffix in ops.DOCUMENT_OP_SUFFIXES:
        assert suffix in FIXTURES, f"{suffix} is advertised with nothing proving it works"
        build, anchor, replacement = FIXTURES[suffix]
        data = build()
        assert ops.outline(data).counts["paragraphs"] > 0, suffix
        edited = ops.replace_text(data, anchor, replacement, path=f"upload{suffix}")
        assert ops.reopens(edited) is None, suffix
        assert replacement in ops.outline(edited).render(), suffix
