"""Synthetic Word documents for the document tests.

Real documents are not committed here — this builds an equivalent one on the
spot, carrying every shape the operations have to survive: heading styles next
to unstyled paragraphs, a bullet list, two tables (one with a merged cell), an
inline picture, a header and a footer with a page-number field, a paragraph
whose runs carry direct formatting, a hyperlink, and Korean/English mixed.
"""

from __future__ import annotations

import io
import struct
import zlib

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


def png_bytes(width: int, height: int, rgb: tuple[int, int, int]) -> bytes:
    """A solid-colour PNG, built without an image library."""
    raw = b"".join(b"\x00" + bytes(rgb) * width for _ in range(height))

    def chunk(kind: bytes, payload: bytes) -> bytes:
        return (
            struct.pack(">I", len(payload))
            + kind
            + payload
            + struct.pack(">I", zlib.crc32(kind + payload) & 0xFFFFFFFF)
        )

    return (
        b"\x89PNG\r\n\x1a\n"
        + chunk(b"IHDR", struct.pack(">IIBBBBB", width, height, 8, 2, 0, 0, 0))
        + chunk(b"IDAT", zlib.compress(raw, 9))
        + chunk(b"IEND", b"")
    )


def _page_number_field(paragraph) -> None:
    """A PAGE field in a footer run — python-docx has no API for one."""
    run = paragraph.add_run()
    begin = run._r.makeelement(qn("w:fldChar"), {qn("w:fldCharType"): "begin"})
    instruction = run._r.makeelement(qn("w:instrText"), {qn("xml:space"): "preserve"})
    instruction.text = " PAGE "
    end = run._r.makeelement(qn("w:fldChar"), {qn("w:fldCharType"): "end"})
    run._r.append(begin)
    run._r.append(instruction)
    run._r.append(end)


def _hyperlink(paragraph, text: str, url: str) -> None:
    """Wrap a run in ``w:hyperlink`` — text the reader sees, in no direct run."""
    from docx.opc.constants import RELATIONSHIP_TYPE as RT

    relationship_id = paragraph.part.relate_to(url, RT.HYPERLINK, is_external=True)
    link = paragraph._p.makeelement(qn("w:hyperlink"), {qn("r:id"): relationship_id})
    run = paragraph._p.makeelement(qn("w:r"), {})
    node = paragraph._p.makeelement(qn("w:t"), {})
    node.text = text
    run.append(node)
    link.append(run)
    paragraph._p.append(link)


def sample_document() -> bytes:
    """The fixture document, as .docx bytes."""
    document = Document()

    section = document.sections[0]
    section.header.paragraphs[0].text = "브레인크루 | Braincrew Confidential"
    footer = section.footer.paragraphs[0]
    footer.text = "page "
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _page_number_field(footer)

    document.add_heading("2026 반영계획안", level=1)              # p0
    document.add_paragraph("")                                     # p1
    document.add_paragraph("본 문서는 현장 점검 결과를 정리한 내부 자료입니다.")  # p2

    document.add_heading("1. 배경 Background", level=2)            # p3
    styled = document.add_paragraph()                              # p4
    head = styled.add_run("검토 결과, ")
    head.font.size = Pt(11)
    warn = styled.add_run("즉시 조치")
    warn.bold = True
    warn.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)
    warn.font.size = Pt(13)
    tail = styled.add_run("가 필요한 항목이 확인되었습니다.")
    tail.font.size = Pt(11)

    document.add_paragraph("세부 항목은 아래 표를 참고한다.", style="List Bullet")   # p5
    document.add_paragraph("담당 부서와 일정은 별도 협의한다.", style="List Bullet")  # p6

    plain = document.add_table(rows=3, cols=3)                     # t0
    plain.style = "Table Grid"
    for row, values in zip(
        plain.rows,
        [
            ["구분", "항목", "비고"],
            ["안전", "소화기 위치 재배치", "9월 1주"],
            ["환경", "폐기물 분리 표지 부착", "9월 2주"],
        ],
        strict=True,
    ):
        for cell, value in zip(row.cells, values, strict=True):
            cell.text = value

    document.add_heading("2. 현장 사진 Photographs", level=2)      # p7
    document.add_picture(io.BytesIO(png_bytes(120, 80, (0x2E, 0x7D, 0x32))), width=Inches(2.4))
    document.add_paragraph("사진 1. 점검 당일 현장")               # p9

    document.add_heading("3. 세부 일정 Schedule", level=3)         # p10
    merged = document.add_table(rows=3, cols=3)                    # t1
    merged.style = "Table Grid"
    merged.cell(0, 0).merge(merged.cell(0, 2)).text = "9월 조치 계획"
    for index, (week, item) in enumerate([("1주", "소화기"), ("2주", "폐기물")], start=1):
        merged.cell(index, 0).text = week
        merged.cell(index, 1).text = item
        merged.cell(index, 2).text = "진행"

    linked = document.add_paragraph("자세한 내용은 ")             # p11
    _hyperlink(linked, "사내 포털", "https://example.com/portal")
    linked.add_run("에서 확인한다.")

    document.add_paragraph("")                                     # p12
    document.add_paragraph(
        "본 자료는 정보 제공 목적으로만 작성되었으며 대외 배포를 금합니다."
    )                                                              # p13

    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()
