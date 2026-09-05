"""Word documents read with addresses and edited in place.

An uploaded ``.docx`` is the truth, not a markdown approximation of it: the
agent reads the real paragraphs, tables and pictures, changes the few it was
asked to change, and every byte it did not touch travels through untouched.
Two halves make that work.

**Addresses.** :func:`outline` renders the document as numbered lines —
``[p0]`` for the first body paragraph, ``[t0]`` for the first table, ``[img0]``
for the first picture — with each paragraph's style name, so the model can see
the heading hierarchy a plain text dump throws away. Indices are for *reading*
and for pointing at something on screen; they shift the moment a paragraph is
inserted, so every write instead takes a **text anchor** copied from that
output. An anchor that matches zero or several paragraphs is refused loudly
(:class:`AnchorError` names the closest paragraph and the first character that
differs) — a silent no-op is how a document comes back unchanged while the
agent reports success.

**Narrow edits.** :func:`replace_text`, :func:`insert_paragraph`,
:func:`remove_paragraph` and :func:`replace_image` are the whole surface. Each
returns new file bytes built by :func:`repack`: the original ZIP with only the
parts the operation actually changed swapped in. Re-serializing a whole Word
package rewrites parts nobody edited — that is how direct formatting multiplies
and fonts get substituted on a round trip — so parts outside the edit are
copied verbatim, entry metadata included.

Text edits reach across ``w:r`` boundaries. Word splits one visible sentence
into arbitrary runs (a spell-check pass alone will do it), so an anchor the
reader can see almost never lines up with a single run; the replacement lands
in the first run it touches and keeps that run's formatting.

Requires ``python-docx`` — installed by the ``office`` extra.
"""

from __future__ import annotations

import io
import re
import zipfile
from dataclasses import dataclass, field
from difflib import SequenceMatcher
from typing import Any

from .converters import UnsafeArchiveError, ensure_archive_within_limits

DOCUMENT_OP_SUFFIXES: tuple[str, ...] = (".docx",)
"""File suffixes these operations handle (lowercase, with the dot)."""

_MAX_AMBIGUOUS_SHOWN = 5
_MAX_STYLES_SHOWN = 24
# Word's own "picture" default; used when a replacement image carries no
# usable pixel dimensions, so a swap can never collapse a shape to zero.
_MIN_EMU = 1


class MissingDocumentDependencyError(RuntimeError):
    """``python-docx`` is not installed, so documents cannot be opened."""


class DocumentOpError(ValueError):
    """An operation was refused. The message is written for the agent."""


class AnchorError(DocumentOpError):
    """An anchor matched zero paragraphs, or more than one."""


class DocumentPartError(DocumentOpError):
    """The document is not a readable Word package."""


# --- addressing ------------------------------------------------------------------


@dataclass(frozen=True)
class Outline:
    """A document rendered as addressed lines, plus the counts above them."""

    lines: list[str]
    counts: dict[str, int]

    def render(self) -> str:
        """The counts header followed by every addressed line."""
        header = " · ".join(f"{key}: {value}" for key, value in self.counts.items())
        return "\n".join([header, *self.lines])


@dataclass
class _Spot:
    """One editable paragraph: where it lives and the runs holding its text.

    ``part`` is the ZIP entry the paragraph is stored in, carried here so an
    edit knows exactly which part to swap without searching for it again.
    """

    label: str
    part: str
    paragraph: Any
    runs: list[Any] = field(default_factory=list)

    @property
    def text(self) -> str:
        return "".join(run.text for run in self.runs)


def _require_docx() -> Any:
    try:
        import docx  # type: ignore[import-untyped]
    except ImportError as exc:  # pragma: no cover - exercised by the message test
        raise MissingDocumentDependencyError(
            "editing .docx needs python-docx — install langchain-canvas[office]"
        ) from exc
    return docx


def _open(data: bytes, *, path: str = "document.docx") -> Any:
    """Parse ``data`` into a python-docx ``Document`` (archive limits first)."""
    docx = _require_docx()
    try:
        ensure_archive_within_limits(data, path=path)
    except UnsafeArchiveError as exc:
        raise DocumentPartError(str(exc)) from exc
    try:
        return docx.Document(io.BytesIO(data))
    except Exception as exc:  # python-docx raises several unrelated types
        raise DocumentPartError(f"{path} is not a readable Word document ({exc})") from exc


def _runs(paragraph: Any) -> list[Any]:
    """Every ``w:r`` under this paragraph, in reading order.

    ``Paragraph.runs`` returns only direct children, which drops the runs
    inside a hyperlink — text the reader plainly sees. Addressing text the
    edit cannot reach is the silent-wrong failure this module exists to
    avoid, so both halves read the same list.
    """
    from docx.text.run import Run  # type: ignore[import-untyped]

    return [Run(element, paragraph) for element in paragraph._p.xpath(".//w:r")]


def _cell_paragraphs(cell: Any, label: str, part: str) -> list[_Spot]:
    spots: list[_Spot] = []
    for index, paragraph in enumerate(cell.paragraphs):
        spots.append(_Spot(f"{label}/p{index}", part, paragraph, _runs(paragraph)))
    for table_index, table in enumerate(cell.tables):
        spots.extend(_table_paragraphs(table, f"{label}/t{table_index}", part))
    return spots


def _table_paragraphs(table: Any, label: str, part: str) -> list[_Spot]:
    spots: list[_Spot] = []
    for row_index, row in enumerate(table.rows):
        for cell_index, cell in enumerate(row.cells):
            spots.extend(_cell_paragraphs(cell, f"{label}/r{row_index}c{cell_index}", part))
    return spots


def _part_name(part: Any) -> str:
    return str(part.partname).lstrip("/")


def _spots(document: Any) -> list[_Spot]:
    """Every paragraph an anchor can reach — body, tables, headers, footers.

    A merged table cell is reported once per grid position by python-docx, so
    the same paragraph would otherwise look like several matches and every
    anchor inside a merged heading would be refused as ambiguous. Paragraphs
    are kept by element identity, first position wins.
    """
    body = _part_name(document.part)
    spots: list[_Spot] = []
    for index, paragraph in enumerate(document.paragraphs):
        spots.append(_Spot(f"p{index}", body, paragraph, _runs(paragraph)))
    for index, table in enumerate(document.tables):
        spots.extend(_table_paragraphs(table, f"t{index}", body))
    many = len(document.sections) > 1
    for number, section in enumerate(document.sections, start=1):
        for name, story in (("header", section.header), ("footer", section.footer)):
            if story.is_linked_to_previous:
                continue
            part = _part_name(story.part)
            label = f"{name}{number}" if many else name
            for index, paragraph in enumerate(story.paragraphs):
                spots.append(_Spot(f"{label}/p{index}", part, paragraph, _runs(paragraph)))
            for index, table in enumerate(story.tables):
                spots.extend(_table_paragraphs(table, f"{label}/t{index}", part))
    seen: set[int] = set()
    unique: list[_Spot] = []
    for spot in spots:
        key = id(spot.paragraph._p)
        if key in seen:
            continue
        seen.add(key)
        unique.append(spot)
    return unique


def _shape_size_in(shape: Any) -> tuple[float, float]:
    return round(shape.width.inches, 2), round(shape.height.inches, 2)


def _cell_fill(cell: Any) -> str | None:
    """The cell's shading as ``#RRGGBB``, or ``None`` when there is nothing to see.

    ``auto`` and explicit white are both the page showing through — a mark for
    them would say "this cell is special" about most of a generated document.
    """
    values = cell._tc.xpath("./w:tcPr/w:shd/@w:fill")
    fill = values[0] if values else None
    if not fill or str(fill).lower() in ("auto", "ffffff"):
        return None
    return f"#{str(fill).upper()}"


def _paragraph_fill(paragraph: Any) -> str | None:
    """The paragraph's own shading as ``#RRGGBB`` — same silence rule as
    :func:`_cell_fill` (auto and white are the page showing through)."""
    values = paragraph._p.xpath("./w:pPr/w:shd/@w:fill")
    fill = values[0] if values else None
    if not fill or str(fill).lower() in ("auto", "ffffff"):
        return None
    return f"#{str(fill).upper()}"


def _cell_text(cell: Any) -> str:
    """One line of cell text, safe inside a pipe row (bars escaped, paragraph
    breaks folded to ``;``)."""
    return "; ".join(
        part for part in (p.strip() for p in cell.text.split("\n")) if part
    ).replace("|", "\\|")


def _table_lines(table: Any) -> list[str]:
    """A table as pipe rows that keep what the page shows.

    ``cell.text`` alone loses the table's *shape*: python-docx reports a
    merged cell once per grid position, so a heading spanning four columns
    read as the same words four times — and a model editing "the second
    Region cell" was editing a merge, not a duplicate. Shading vanished
    entirely, and a comma inside a cell broke the old comma-joined row.

    Here each merged cell appears once, at its top-left, sized in the reading
    order a person uses: ``(2x3)`` is two columns wide, three rows tall. A
    vertical continuation shows as ``^`` so columns keep their place. A cell's
    shading rides it as ``[#RRGGBB]``, a nested table as ``[+N nested]``, and
    a row marked to repeat on every page as ``(header row)``.
    """
    # One walk of the grid, holding on to the cell objects: lxml re-creates
    # element proxies on re-access, so an ``id()`` taken in a first pass does
    # not survive a second ``row.cells`` — the kept references are what make
    # the merge detection stable.
    all_rows = list(table.rows)
    grid = [list(row.cells) for row in all_rows]
    origin: dict[int, tuple[int, int]] = {}
    span: dict[int, list[int]] = {}  # id(tc) -> [colspan, rowspan]
    for row_index, row_cells in enumerate(grid):
        for col_index, cell in enumerate(row_cells):
            tc = id(cell._tc)
            if tc not in origin:
                origin[tc] = (row_index, col_index)
                span[tc] = [1, 1]
            else:
                first_row, first_col = origin[tc]
                span[tc][0] = max(span[tc][0], col_index - first_col + 1)
                span[tc][1] = max(span[tc][1], row_index - first_row + 1)

    lines: list[str] = []
    for row_index, row in enumerate(all_rows):
        cells: list[str] = []
        emitted_here: set[int] = set()
        for cell in grid[row_index]:
            tc = id(cell._tc)
            if tc in emitted_here:
                continue  # horizontal continuation — the (CxR) already says how wide
            emitted_here.add(tc)
            if origin[tc][0] != row_index:
                cells.append("^")  # vertical continuation — keeps columns in place
                continue
            parts = [_cell_text(cell)]
            colspan, rowspan = span[tc]
            if colspan > 1 or rowspan > 1:
                parts.append(f"({colspan}x{rowspan})")
            fill = _cell_fill(cell)
            if fill:
                parts.append(f"[{fill}]")
            if cell.tables:
                parts.append(f"[+{len(cell.tables)} nested]")
            cells.append(" ".join(part for part in parts if part))
        line = "      | " + " | ".join(cells) + " |"
        if row._tr.xpath("./w:trPr/w:tblHeader"):
            line += " (header row)"
        lines.append(line)
    return lines


def outline(data: bytes, *, path: str = "document.docx") -> Outline:
    """The document as addressed lines: paragraphs, tables and pictures.

    Body paragraphs are numbered by their position in the document, blanks
    included, so ``[p7]`` is ``document.paragraphs[7]`` for any reader —
    the on-screen preview included. Empty paragraphs are numbered but not
    printed; the gap in the numbering is the honest record of them.
    """
    document = _open(data, path=path)
    from docx.table import Table  # type: ignore[import-untyped]

    shape_size: dict[int, tuple[float, float]] = {}
    for index, shape in enumerate(document.inline_shapes):
        shape_size[index] = _shape_size_in(shape)

    lines: list[str] = []
    paragraph_index = 0
    table_index = 0
    image_index = 0
    for item in document.iter_inner_content():
        if isinstance(item, Table):
            lines.append(f"[t{table_index}] {len(item.rows)}x{len(item.columns)} table")
            lines.extend(_table_lines(item))
            table_index += 1
            continue
        text = item.text.strip()
        style = item.style.name if item.style is not None else ""
        prefix = f"[p{paragraph_index}]"
        if text:
            named = style and style != "Normal"
            line = f"{prefix} ({style}) {text}" if named else f"{prefix} {text}"
            # A paragraph's own background (a highlight box) is part of what
            # the page shows; same silence rule as a table cell's shading.
            fill = _paragraph_fill(item)
            lines.append(f"{line} [{fill}]" if fill else line)
        for _ in item._p.xpath(".//w:drawing//a:blip"):
            width, height = shape_size.get(image_index, (0.0, 0.0))
            lines.append(f"[img{image_index}] {width} x {height} in, in p{paragraph_index}")
            image_index += 1
        paragraph_index += 1

    for spot in _spots(document):
        if not spot.label.startswith(("header", "footer")):
            continue
        text = spot.text.strip()
        if text:
            lines.append(f"[{spot.label}] {text}")

    counts = {
        "paragraphs": len(document.paragraphs),
        "tables": len(document.tables),
        "images": len(document.inline_shapes),
        "sections": len(document.sections),
    }
    return Outline(lines=lines, counts=counts)


# --- anchors ---------------------------------------------------------------------


def _find(spots: list[_Spot], anchor: str) -> list[tuple[_Spot, int]]:
    """Every ``(spot, offset)`` where ``anchor`` occurs, non-overlapping."""
    hits: list[tuple[_Spot, int]] = []
    for spot in spots:
        text = spot.text
        start = text.find(anchor)
        while start != -1:
            hits.append((spot, start))
            start = text.find(anchor, start + len(anchor))
    return hits


def _first_difference(anchor: str, candidate: str) -> str:
    """Where the anchor and the closest paragraph part ways, in words."""
    limit = min(len(anchor), len(candidate))
    for index in range(limit):
        if anchor[index] != candidate[index]:
            return (
                f"First difference at character {index + 1}: you wrote "
                f"{anchor[index]!r}, the document has {candidate[index]!r}."
            )
    # The other way round cannot reach here: an anchor that is a prefix of a
    # paragraph is found, so it never becomes a near miss.
    return (
        f"The document's paragraph ends after {len(candidate)} characters; "
        f"your anchor runs {len(anchor) - len(candidate)} character(s) longer."
    )


def _not_found(spots: list[_Spot], anchor: str) -> AnchorError:
    """A refusal that says which paragraph was nearly it, and how it differed."""
    best: _Spot | None = None
    best_ratio = 0.0
    for spot in spots:
        text = spot.text
        if not text.strip():
            continue
        ratio = SequenceMatcher(None, anchor, text).ratio()
        if ratio > best_ratio:
            best, best_ratio = spot, ratio
    parts = [f"anchor not found (0 matches): {anchor!r}."]
    if "\n" in anchor:
        parts.append("An anchor has to sit inside one paragraph — this one spans lines.")
    if best is not None:
        parts.append(f"Closest paragraph [{best.label}]: {best.text!r}")
        parts.append("  " + _first_difference(anchor, best.text))
    parts.append(
        "  Copy the anchor from the read_canvas output rather than retyping it; "
        "one character apart matches nothing."
    )
    return AnchorError("\n".join(parts))


# An anchor may lead with the address the reader printed: ``[p7] 리서치``.
_ADDRESSED = re.compile(r"^\[([A-Za-z0-9/]+)\]\s*(.*)$", re.DOTALL)


def split_address(anchor: str) -> tuple[str | None, str]:
    """``("p7", "리서치")`` for ``"[p7] 리서치"``; ``(None, anchor)`` otherwise.

    A title that is the whole of its paragraph and appears again in the body
    cannot be made unique by "extending it with surrounding words" — there
    are none. The address the reader prints is the disambiguator, and the
    one a model reaches for first; so the anchor accepts it in front. The
    text after it may be empty, which means the whole paragraph.
    """
    match = _ADDRESSED.match(anchor)
    if match is None:
        return None, anchor
    return match.group(1), match.group(2)


def _resolve(spots: list[_Spot], anchor: str) -> tuple[_Spot, int]:
    """The single paragraph an anchor points at, or a loud refusal.

    Returns ``(spot, offset)``; ``offset`` is where the anchor's text part
    starts inside the paragraph (0 for an address-only anchor).
    """
    label, text = split_address(anchor)
    if label is not None:
        addressed = [spot for spot in spots if spot.label == label]
        if not addressed:
            known = ", ".join(f"[{spot.label}]" for spot in spots[:_MAX_AMBIGUOUS_SHOWN])
            raise AnchorError(
                f"no paragraph is addressed [{label}] — the addresses in this document "
                f"start {known}; copy one from read_canvas."
            )
        spots = addressed
        if not text:
            return spots[0], 0
        hits = _find(spots, text)
        if not hits:
            raise AnchorError(
                f"[{label}] does not contain {text!r} — that paragraph reads "
                f"{spots[0].text!r}."
            )
        return hits[0]
    if not anchor:
        raise AnchorError("anchor is empty — pass the text you want to point at.")
    hits = _find(spots, anchor)
    if not hits:
        raise _not_found(spots, anchor)
    if len(hits) > 1:
        where = ", ".join(f"[{spot.label}]" for spot, _ in hits[:_MAX_AMBIGUOUS_SHOWN])
        extra = len(hits) - _MAX_AMBIGUOUS_SHOWN
        more = f" and {extra} more" if extra > 0 else ""
        raise AnchorError(
            f"anchor matches {len(hits)} places ({where}{more}) — put the address in "
            f"front to pick one, like {where.split(', ')[0]} {anchor!r}."
        )
    return hits[0]


# --- repacking -------------------------------------------------------------------


def repack(
    original: bytes,
    saved: bytes,
    changed: set[str],
    *,
    removed: set[str] | None = None,
) -> bytes:
    """``original`` with only ``changed`` (and brand-new) entries from ``saved``.

    Every other entry — styles, theme, fonts, numbering, media the edit never
    looked at — is copied with its bytes and its ZIP metadata intact, so a
    checksum of any untouched part matches the file the user uploaded.

    ``removed`` names entries to leave out. An operation has to say so
    explicitly: dropping a part is a change like any other, and a package
    that quietly loses whatever the edit did not mention is how a document
    comes back missing a picture nobody meant to touch.
    """
    gone = removed or set()
    with zipfile.ZipFile(io.BytesIO(original)) as source, zipfile.ZipFile(
        io.BytesIO(saved)
    ) as edited:
        source_names = source.namelist()
        known = set(source_names)
        out = io.BytesIO()
        with zipfile.ZipFile(out, "w", zipfile.ZIP_DEFLATED) as result:
            for info in source.infolist():
                if info.filename in gone:
                    continue
                take_edited = info.filename in changed and info.filename in edited.namelist()
                payload = edited.read(info.filename) if take_edited else source.read(info.filename)
                result.writestr(
                    zipfile.ZipInfo(info.filename, date_time=info.date_time),
                    payload,
                    compress_type=info.compress_type,
                )
            for info in edited.infolist():
                if info.filename in known or info.filename in gone:
                    continue
                result.writestr(info, edited.read(info.filename))
        return out.getvalue()


def _saved(document: Any, *, drop_unused: str | None = None) -> bytes:
    """The package as python-docx writes it, optionally letting one image go.

    ``drop_unused`` is a relationship id the caller has just stopped using.
    It is released only when nothing in that part points at it any more — the
    same picture can appear twice under one id, and one of them may still be
    on the page. Releasing it is what lets python-docx's own walk of the
    package decide whether the image behind it is still reachable: a picture
    the header also shows stays, because the header's own relationship still
    reaches it.
    """
    if drop_unused is not None and _relationship_uses(document.part, drop_unused) == 0:
        document.part.drop_rel(drop_unused)
    out = io.BytesIO()
    document.save(out)
    return out.getvalue()


def _relationship_uses(part: Any, relationship_id: str) -> int:
    """How many times this part's XML still points at ``relationship_id``."""
    used = part._element.xpath("//@r:embed | //@r:link")
    return sum(1 for value in used if value == relationship_id)


def _orphaned(original: bytes, saved: bytes) -> set[str]:
    """Media entries the original carried that the saved package dropped.

    Only ``word/media/`` — that is where an image lives, and it is the one
    place a replaced picture leaves bytes behind. Anything else missing from
    a re-serialization would be a surprise, not a result, so it is kept.
    """
    with zipfile.ZipFile(io.BytesIO(original)) as before, zipfile.ZipFile(
        io.BytesIO(saved)
    ) as after:
        still_there = set(after.namelist())
        return {
            name
            for name in before.namelist()
            if name.startswith("word/media/") and name not in still_there
        }




# --- operations ------------------------------------------------------------------


def replace_text(data: bytes, old: str, new: str, *, path: str = "document.docx") -> bytes:
    """Replace the one occurrence of ``old`` with ``new``, runs crossed.

    Searches body paragraphs, table cells (nested tables included), headers
    and footers. Raises :class:`AnchorError` unless exactly one place matches.
    """
    document = _open(data, path=path)
    spots = _spots(document)
    spot, start = _resolve(spots, old)
    label, matched = split_address(old)
    if label is not None:
        # The address is how the caller pointed, not text for the page — a
        # replacement written in the same shape ("[p7] new title") drops it.
        new_label, new_text = split_address(new)
        if new_label == label:
            new = new_text
    length = len(matched) if matched else len(spot.text)
    _splice(spot.runs, start, start + length, new)
    return repack(data, _saved(document), {spot.part})


def _splice(runs: list[Any], start: int, end: int, new: str) -> None:
    """Put ``new`` in place of ``[start, end)`` across the runs it spans.

    The text lands in the first run it touches, so it inherits that run's
    formatting — the same thing typing over a selection does in Word.
    """
    position = 0
    injected = False
    for run in runs:
        low, high = position, position + len(run.text)
        position = high
        if high <= start or low >= end:
            continue
        head = run.text[: max(0, min(start - low, high - low))]
        tail = run.text[max(0, min(end - low, high - low)) :]
        run.text = head + ("" if injected else new) + tail
        injected = True


def insert_paragraph(
    data: bytes,
    *,
    anchor: str,
    text: str,
    style: str | None = None,
    position: str = "after",
    path: str = "document.docx",
) -> bytes:
    """Add a paragraph next to the paragraph ``anchor`` points at."""
    if position not in {"after", "before"}:
        raise DocumentOpError(f'position is "after" or "before" (got {position!r}).')
    document = _open(data, path=path)
    spots = _spots(document)
    spot, _ = _resolve(spots, anchor)
    try:
        added = spot.paragraph.insert_paragraph_before(text, style)
    except KeyError as exc:
        raise DocumentOpError(_unknown_style(document, style)) from exc
    if position == "after":
        spot.paragraph._p.addnext(added._p)
    return repack(data, _saved(document), {spot.part})


def _unknown_style(document: Any, style: str | None) -> str:
    from docx.enum.style import WD_STYLE_TYPE  # type: ignore[import-untyped]

    names = sorted(
        item.name
        for item in document.styles
        if item.type == WD_STYLE_TYPE.PARAGRAPH and item.name
    )
    shown = ", ".join(names[:_MAX_STYLES_SHOWN])
    more = "" if len(names) <= _MAX_STYLES_SHOWN else f", and {len(names) - _MAX_STYLES_SHOWN} more"
    return (
        f"this document has no paragraph style named {style!r}. Styles it does "
        f"have: {shown}{more}. Omit `style` to use the default one."
    )


def remove_paragraph(data: bytes, *, anchor: str, path: str = "document.docx") -> bytes:
    """Delete the paragraph ``anchor`` points at."""
    document = _open(data, path=path)
    spots = _spots(document)
    spot, _ = _resolve(spots, anchor)
    element = spot.paragraph._p
    parent = element.getparent()
    siblings = parent.findall(element.tag)
    if len(siblings) <= 1:
        raise DocumentOpError(
            f"[{spot.label}] is the only paragraph in its container; Word needs one "
            "there. Replace its text instead of removing it."
        )
    parent.remove(element)
    return repack(data, _saved(document), {spot.part})


def _body_ancestor(document: Any, element: Any) -> Any | None:
    """The top-level body child ``element`` sits under, or None if it is elsewhere."""
    body = document.element.body
    node = element
    while node is not None and node.getparent() is not body:
        node = node.getparent()
    return node


def _section_of(document: Any, element: Any) -> Any:
    """The section that governs where ``element`` sits.

    A section ends at the paragraph carrying its ``w:sectPr``, so the number of
    those before a body element is that element's section. A header, a footer
    and the end of the document all belong to the last section.
    """
    sections = document.sections
    top = _body_ancestor(document, element)
    if top is None:
        return sections[-1]
    index = 0
    for child in document.element.body.iterchildren():
        if child is top:
            break
        if child.xpath("./w:pPr/w:sectPr"):
            index += 1
    return sections[min(index, len(sections) - 1)]


def _text_width(section: Any) -> int:
    """How wide the text column is in this section, in EMU."""
    page = section.page_width
    left = section.left_margin
    right = section.right_margin
    if page is None or left is None or right is None:
        return 0
    return max(_MIN_EMU, int(page) - int(left) - int(right))


def _rels_name(part: str) -> str:
    """Where a part's relationships live, given the part's own name."""
    folder, _, name = part.rpartition("/")
    return f"{folder}/_rels/{name}.rels" if folder else f"_rels/{name}.rels"


def insert_image(
    data: bytes,
    *,
    image: bytes,
    anchor: str | None = None,
    position: str = "after",
    width_inches: float | None = None,
    alt_text: str | None = None,
    path: str = "document.docx",
) -> tuple[bytes, str]:
    """Add ``image`` to the document as a paragraph of its own.

    Without an ``anchor`` the picture goes at the end of the document, which
    is where a reader asking for "a picture at the bottom" means. With one it
    goes next to the paragraph that anchor names, the same way a paragraph is
    inserted.

    The default width is the image's own, brought down to the width of the
    text column when it is wider than that — never up. A stated
    ``width_inches`` is used as given and the height follows the image's
    proportions, the rule a replacement already follows.

    Returns the new file bytes and a one-line note about the size.
    """
    from docx.image.image import Image  # type: ignore[import-untyped]
    from docx.shared import Emu, Inches  # type: ignore[import-untyped]

    if position not in {"after", "before"}:
        raise DocumentOpError(f'position is "after" or "before" (got {position!r}).')
    if width_inches is not None and width_inches <= 0:
        raise DocumentOpError(f"width_inches has to be more than zero (got {width_inches}).")
    try:
        picture = Image.from_blob(image)
    except Exception as exc:  # python-docx raises several unrelated types
        raise DocumentOpError(f"that file is not a readable image ({exc}).") from exc
    document = _open(data, path=path)
    if anchor is None:
        paragraph = document.add_paragraph()
        part = _part_name(document.part)
    else:
        spot, _ = _resolve(_spots(document), anchor)
        paragraph = spot.paragraph.insert_paragraph_before()
        if position == "after":
            spot.paragraph._p.addnext(paragraph._p)
        part = spot.part
    room = _text_width(_section_of(document, paragraph._p))
    natural = int(picture.width)
    if width_inches is None:
        width = min(natural, room) if room else natural
    else:
        width = int(Inches(width_inches))
    ratio = picture.px_height / picture.px_width if picture.px_width else 1.0
    height = max(_MIN_EMU, int(round(width * ratio)))
    run = paragraph.add_run()
    shape = run.add_picture(io.BytesIO(image), width=Emu(width), height=Emu(height))
    if alt_text:
        # python-docx has no API for alt text; the attribute is where Word
        # reads it from, and a screen reader after it.
        shape._inline.docPr.set("descr", alt_text)
    changed = {part, _rels_name(part), "[Content_Types].xml"}
    shrunk = width_inches is None and width < natural
    size = f"{round(shape.width.inches, 2)} x {round(shape.height.inches, 2)} in"
    note = f"The picture is {size}" + (
        " — scaled down to the width of the text column." if shrunk else "."
    )
    return repack(data, _saved(document), changed), note


def replace_image(
    data: bytes, *, index: int, image: bytes, path: str = "document.docx"
) -> tuple[bytes, str]:
    """Swap the bytes behind ``[img<index>]``, keeping its width on the page.

    The height is refitted to the replacement's own proportions, so the
    picture keeps the column width the document was laid out around instead
    of stretching. Returns the new file bytes and a one-line note about the
    size, for the caller to relay.
    """
    from docx.image.image import Image  # type: ignore[import-untyped]
    from docx.shared import Emu  # type: ignore[import-untyped]

    document = _open(data, path=path)
    shapes = document.inline_shapes
    if not shapes:
        raise DocumentOpError("this document has no pictures to replace.")
    if index < 0 or index >= len(shapes):
        raise DocumentOpError(
            f"[img{index}] is not in this document — it has {len(shapes)} picture(s), "
            f"addressed [img0]..[img{len(shapes) - 1}]."
        )
    try:
        replacement = Image.from_blob(image)
    except Exception as exc:
        raise DocumentOpError(f"the replacement is not a readable image ({exc}).") from exc
    shape = shapes[index]
    old_width, old_height = _shape_size_in(shape)
    blip = shape._inline.graphic.graphicData.pic.blipFill.blip
    previous_id = blip.embed
    relationship_id, _ = document.part.get_or_add_image(io.BytesIO(image))
    blip.embed = relationship_id
    width = shape.width.emu
    ratio = replacement.px_height / replacement.px_width if replacement.px_width else 1.0
    shape.height = Emu(max(_MIN_EMU, int(round(width * ratio))))
    new_width, new_height = _shape_size_in(shape)
    changed = {
        _part_name(document.part),
        "word/_rels/document.xml.rels",
        "[Content_Types].xml",
    }
    saved = _saved(document, drop_unused=previous_id)
    dropped = _orphaned(data, saved)
    note = (
        f"[img{index}] was {old_width} x {old_height} in and is now "
        f"{new_width} x {new_height} in (width kept, height refitted)."
    )
    return repack(data, saved, changed, removed=dropped), note


# --- verification ----------------------------------------------------------------


def reopens(data: bytes) -> str | None:
    """``None`` when the bytes parse as a Word document, else why they do not."""
    try:
        _open(data)
    except (DocumentPartError, MissingDocumentDependencyError) as exc:
        return str(exc)
    return None


# --- table structure & look -------------------------------------------------------

_TABLE_LABEL = re.compile(r"^\[?t(\d+)\]?$")
_CELL_RANGE = re.compile(r"^r(\d+)c(\d+)(?:\s*:\s*r(\d+)c(\d+))?$")
_FILL_HEX = re.compile(r"^#?([0-9a-fA-F]{6})$")


def _body_table(document: Any, table: str) -> Any:
    """The body table ``"t0"`` / ``"[t0]"`` names, refused loudly otherwise.

    These are the tables :func:`outline` numbers; a table living in a header
    or footer has no ``[tN]`` address and is out of reach here on purpose.
    """
    match = _TABLE_LABEL.match(table.strip())
    if match is None:
        raise DocumentOpError(
            f'{table!r} is not a table address — outline names tables "[t0]", "[t1]", ...'
        )
    index = int(match.group(1))
    tables = document.tables
    if index >= len(tables):
        have = f"{len(tables)} table(s)" if tables else "no tables"
        raise DocumentOpError(f"this document has {have} — [t{index}] does not exist")
    return tables[index]


def _cell_rectangle(table_obj: Any, cells: str) -> tuple[int, int, int, int]:
    """``"r0c0"`` or ``"r0c0:r1c2"`` as a normalized in-bounds rectangle."""
    match = _CELL_RANGE.match(cells.strip())
    if match is None:
        raise DocumentOpError(
            f'{cells!r} is not a cell address — write "r0c0" for one cell or '
            '"r0c0:r1c2" for a rectangle (0-based, from outline\'s grid)'
        )
    r1, c1 = int(match.group(1)), int(match.group(2))
    r2 = int(match.group(3)) if match.group(3) is not None else r1
    c2 = int(match.group(4)) if match.group(4) is not None else c1
    top, bottom = min(r1, r2), max(r1, r2)
    left, right = min(c1, c2), max(c1, c2)
    n_rows, n_cols = len(table_obj.rows), len(table_obj.columns)
    if bottom >= n_rows or right >= n_cols:
        raise DocumentOpError(
            f"{cells} reaches past the table — it is {n_rows}x{n_cols} "
            f"(rows r0-r{n_rows - 1}, columns c0-c{n_cols - 1})"
        )
    return top, left, bottom, right


def merge_table_cells(
    data: bytes, table: str, cells: str, *, path: str = "document.docx"
) -> bytes:
    """Merge the rectangle ``cells`` of body table ``table`` into one cell.

    ``cells`` is ``"r0c0:r0c2"`` — the 0-based grid the outline prints. The
    merged cell keeps every member's text (Word appends the paragraphs, the
    same thing merging in Word itself does); prune the extra words with a
    text edit afterwards if the copies should go. A rectangle that cuts
    across an existing merge is refused rather than guessed at.
    """
    document = _open(data, path=path)
    table_obj = _body_table(document, table)
    top, left, bottom, right = _cell_rectangle(table_obj, cells)
    if (top, left) == (bottom, right):
        raise DocumentOpError(
            f'{cells} is a single cell — a merge needs a rectangle, like "r0c0:r0c2"'
        )
    try:
        table_obj.cell(top, left).merge(table_obj.cell(bottom, right))
    except Exception as exc:  # python-docx raises its own span errors
        raise DocumentOpError(
            f"cannot merge {cells}: {exc} — a merge must be a clean rectangle "
            "and may not cut across an existing merge"
        ) from exc
    return repack(data, _saved(document), {_part_name(document.part)})


def style_table_cells(
    data: bytes,
    table: str,
    cells: str,
    *,
    fill: str | None = None,
    align: str | None = None,
    path: str = "document.docx",
) -> bytes:
    """Restyle the cells of body table ``table`` in the rectangle ``cells``.

    ``fill`` is ``"#RRGGBB"`` shading, or ``"none"`` to take existing shading
    off; ``align`` is ``left`` / ``center`` / ``right`` for every paragraph in
    the cell. At least one must be given — a call that changes nothing is a
    mistake worth naming. A merged cell inside the rectangle is touched once.
    """
    if fill is None and align is None:
        raise DocumentOpError("give fill (#RRGGBB or none), align (left/center/right), or both")
    clear_fill = isinstance(fill, str) and fill.strip().lower() == "none"
    fill_hex: str | None = None
    if fill is not None and not clear_fill:
        match = _FILL_HEX.match(fill.strip())
        if match is None:
            raise DocumentOpError(f'{fill!r} is not a fill — write "#RRGGBB", or "none" to clear')
        fill_hex = match.group(1).upper()
    if align is not None and align not in ("left", "center", "right"):
        raise DocumentOpError(f"{align!r} is not an alignment — left, center or right")

    document = _open(data, path=path)
    table_obj = _body_table(document, table)
    top, left, bottom, right = _cell_rectangle(table_obj, cells)

    from docx.enum.text import WD_ALIGN_PARAGRAPH  # type: ignore[import-untyped]
    from docx.oxml import OxmlElement  # type: ignore[import-untyped]
    from docx.oxml.ns import qn  # type: ignore[import-untyped]

    alignments = {
        "left": WD_ALIGN_PARAGRAPH.LEFT,
        "center": WD_ALIGN_PARAGRAPH.CENTER,
        "right": WD_ALIGN_PARAGRAPH.RIGHT,
    }
    touched: set[int] = set()  # a merged cell repeats per grid position; once is enough
    cells_in_range = [
        table_obj.cell(row, column)
        for row in range(top, bottom + 1)
        for column in range(left, right + 1)
    ]
    for cell in cells_in_range:
        marker = id(cell._tc)
        if marker in touched:
            continue
        touched.add(marker)
        if fill is not None:
            properties = cell._tc.get_or_add_tcPr()
            for shading in properties.findall(qn("w:shd")):
                properties.remove(shading)
            if fill_hex is not None:
                shading = OxmlElement("w:shd")
                shading.set(qn("w:fill"), fill_hex)
                properties.append(shading)
        if align is not None:
            for paragraph in cell.paragraphs:
                paragraph.alignment = alignments[align]
    return repack(data, _saved(document), {_part_name(document.part)})
